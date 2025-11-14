"""
V6.1 Extractor - Better education detection + improved Word
- More education keywords
- Even better Word handling
- Better section detection
"""
import pdfplumber
import PyPDF2
from docx import Document as DocxDocument
import re
from typing import Dict, List, Tuple, Optional
from io import BytesIO


class DocumentExtractor:
    """V6.1 with improved detection"""

    def __init__(self):
        self.bullet_markers = ['•', '-', '●', '○', '*', '»', '→', '▪', '▫', '–', '—', '·', '►', '➤']

    def extract_from_pdf(self, file_bytes: bytes) -> Dict:
        """Extract from PDF - Working well"""
        try:
            pdf_file = BytesIO(file_bytes)
            all_lines = []

            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        all_lines.extend(page_text.split('\n'))

            merged_lines = self._merge_continuation_lines(all_lines)
            full_text = '\n'.join(merged_lines)
            sections = self._parse_sections_complete(merged_lines)

            return {
                'success': True,
                'text': full_text,
                'sections': sections,
                'lines': merged_lines,
                'word_count': len(full_text.split()),
                'line_count': len([l for l in merged_lines if l.strip()])
            }

        except Exception as e:
            return {
                'success': False,
                'error': f"PDF extraction failed: {str(e)}"
            }

    def extract_from_word(self, file_bytes: bytes) -> Dict:
        """
        IMPROVED Word extraction - Even more careful
        """
        try:
            doc = DocxDocument(BytesIO(file_bytes))
            lines = []

            # Extract paragraphs - Word has good structure
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Check if it's formatted (bold = likely header or title)
                    is_bold = False
                    if paragraph.runs:
                        bold_count = sum(1 for run in paragraph.runs if run.bold)
                        is_bold = bold_count > len(paragraph.runs) / 2

                    # If bold and short, likely a header
                    if is_bold and len(text.split()) <= 5:
                        # Mark as potential header by making it caps
                        if not text.isupper():
                            text = text.upper()

                    lines.append(text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    cells_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and cell_text not in cells_text:  # Avoid duplicates
                            cells_text.append(cell_text)
                    if cells_text:
                        lines.append(' | '.join(cells_text))

            # Very gentle merging for Word (it's already structured)
            merged_lines = []
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    i += 1
                    continue

                # Check if section header
                if self._is_section_header_strict(line):
                    merged_lines.append(line)
                    i += 1
                    continue

                # Check if bullet
                if self._is_bullet_start(line):
                    merged_lines.append(line)
                    i += 1
                    continue

                # Otherwise, check if next line is continuation
                current_text = line
                j = i + 1

                # Only merge if next line is short and lowercase
                while j < len(lines):
                    next_line = lines[j].strip()
                    if not next_line:
                        j += 1
                        continue

                    # Stop if it's a header or bullet
                    if self._is_section_header_strict(next_line) or self._is_bullet_start(next_line):
                        break

                    # Only merge if clearly continuation (lowercase start, short)
                    if len(next_line) < 40 and next_line[0].islower():
                        current_text += ' ' + next_line
                        j += 1
                    else:
                        break

                merged_lines.append(current_text)
                i = j if j > i + 1 else i + 1

            full_text = '\n'.join(merged_lines)
            sections = self._parse_sections_complete(merged_lines)

            return {
                'success': True,
                'text': full_text,
                'sections': sections,
                'lines': merged_lines,
                'word_count': len(full_text.split()),
                'line_count': len(merged_lines)
            }

        except Exception as e:
            return {
                'success': False,
                'error': f"Word extraction failed: {str(e)}"
            }

    def _merge_continuation_lines(self, lines: List[str]) -> List[str]:
        """PDF merging - Works well"""
        merged = []
        current_bullet = None
        current_paragraph = None

        for i, line in enumerate(lines):
            stripped = line.strip()

            if not stripped:
                if current_bullet:
                    merged.append(current_bullet)
                    current_bullet = None
                elif current_paragraph:
                    merged.append(current_paragraph)
                    current_paragraph = None
                continue

            is_bullet_start = self._is_bullet_start(stripped)
            is_header = self._is_section_header_strict(stripped)

            if is_header:
                if current_bullet:
                    merged.append(current_bullet)
                    current_bullet = None
                if current_paragraph:
                    merged.append(current_paragraph)
                    current_paragraph = None
                merged.append(line)
                continue

            if is_bullet_start:
                if current_bullet:
                    merged.append(current_bullet)
                if current_paragraph:
                    merged.append(current_paragraph)
                    current_paragraph = None
                current_bullet = line
                continue

            is_continuation = self._is_continuation(stripped, lines, i)

            if is_continuation:
                if current_bullet:
                    current_bullet += ' ' + stripped
                elif current_paragraph:
                    current_paragraph += ' ' + stripped
                else:
                    current_paragraph = line
            else:
                if current_bullet:
                    merged.append(current_bullet)
                    current_bullet = None
                if current_paragraph:
                    merged.append(current_paragraph)
                    current_paragraph = None

                if not is_bullet_start and len(stripped) > 20:
                    current_paragraph = line
                else:
                    merged.append(line)

        if current_bullet:
            merged.append(current_bullet)
        if current_paragraph:
            merged.append(current_paragraph)

        return merged

    def _is_bullet_start(self, line: str) -> bool:
        """Check if line starts a bullet"""
        if not line:
            return False

        # Bullet symbols
        if any(line.startswith(marker) for marker in self.bullet_markers):
            return True

        # Numbered bullets
        if re.match(r'^\d+[\.)]\s', line):
            return True

        # Action verbs
        if len(line) > 40:
            action_verbs = [
                'achieved', 'developed', 'created', 'managed', 'led', 'implemented',
                'designed', 'built', 'improved', 'launched', 'delivered', 'collaborated',
                'spearheaded', 'established', 'optimized', 'coordinated', 'drove',
                'executed', 'engineered', 'architected', 'streamlined', 'reduced',
                'increased', 'generated', 'transformed', 'automated', 'analyzed'
            ]

            first_word = line.split()[0].lower() if line else ""
            if first_word in action_verbs and line[0].isupper():
                return True

        return False

    def _is_section_header_strict(self, line: str) -> bool:
        """
        STRICT section header detection
        IMPROVED: More education keywords
        """
        if not line:
            return False

        line_upper = line.upper()
        line_stripped = line.strip()

        # Expanded section keywords
        known_sections = [
            # Summary
            'SUMMARY', 'PROFESSIONAL SUMMARY', 'PROFILE', 'OBJECTIVE', 'ABOUT ME', 'OVERVIEW',
            # Experience
            'EXPERIENCE', 'WORK EXPERIENCE', 'PROFESSIONAL EXPERIENCE', 'EMPLOYMENT',
            'WORK HISTORY', 'CAREER HISTORY', 'EMPLOYMENT HISTORY',
            # Education - EXPANDED
            'EDUCATION', 'ACADEMIC BACKGROUND', 'ACADEMIC QUALIFICATIONS', 'EDUCATIONAL BACKGROUND',
            'ACADEMIC CREDENTIALS', 'DEGREES', 'ACADEMIC HISTORY', 'EDUCATIONAL QUALIFICATIONS',
            'QUALIFICATIONS', 'TRAINING', 'ACADEMIC TRAINING',
            # Skills
            'SKILLS', 'TECHNICAL SKILLS', 'CORE COMPETENCIES', 'TECHNOLOGIES', 'EXPERTISE',
            'PROFICIENCIES', 'TECHNICAL PROFICIENCIES', 'TECHNICAL EXPERTISE',
            # Other sections
            'PROJECTS', 'KEY PROJECTS', 'NOTABLE PROJECTS', 'PROFESSIONAL PROJECTS',
            'CERTIFICATIONS', 'CERTIFICATES', 'LICENSES', 'CREDENTIALS', 'PROFESSIONAL CERTIFICATIONS',
            'AWARDS', 'HONORS', 'ACHIEVEMENTS', 'RECOGNITION', 'ACCOMPLISHMENTS', 'HONORS AND AWARDS',
            'PUBLICATIONS', 'RESEARCH', 'PAPERS', 'RESEARCH PUBLICATIONS',
            'VOLUNTEER', 'VOLUNTEER EXPERIENCE', 'COMMUNITY SERVICE', 'VOLUNTEER WORK',
            'LANGUAGES', 'LANGUAGE SKILLS', 'LANGUAGE PROFICIENCY',
            'INTERESTS', 'HOBBIES', 'ACTIVITIES', 'PERSONAL INTERESTS'
        ]

        # Exact match (case insensitive)
        for section in known_sections:
            if line_upper == section or line_upper == section + ':':
                return True

        # Check partial matches for common patterns
        # "EDUCATION & TRAINING", "SKILLS & EXPERTISE", etc.
        for section in known_sections:
            if section in line_upper and len(line_upper) < 40:
                # Make sure it's not part of a longer sentence
                if line_upper.count(' ') <= 5:
                    return True

        # All caps and very short (2-4 words)
        word_count = len(line_stripped.split())
        if line_stripped == line_upper and 1 <= word_count <= 4:
            # Must not start with bullet
            if any(line_stripped.startswith(m) for m in self.bullet_markers):
                return False
            # Must not have numbers (dates, etc)
            if re.search(r'\d', line_stripped):
                return False
            # Must be reasonable length
            if 3 < len(line_stripped) < 35:
                return True

        return False

    def _is_continuation(self, line: str, all_lines: List[str], current_idx: int) -> bool:
        """Check if line is continuation"""
        if not line:
            return False

        if line[0].islower():
            return True

        if len(line) < 50:
            prev_idx = current_idx - 1
            while prev_idx >= 0 and not all_lines[prev_idx].strip():
                prev_idx -= 1

            if prev_idx >= 0:
                prev_line = all_lines[prev_idx].strip()
                if prev_line and not prev_line.endswith(('.', '!', '?', ':')):
                    return True

        return False

    def _parse_sections_complete(self, lines: List[str]) -> Dict:
        """Parse ALL sections"""
        sections = {
            'header': {'content': '', 'lines': [], 'blocks': []},
            'summary': {'content': '', 'lines': [], 'blocks': []},
            'experience': {'content': '', 'lines': [], 'jobs': []},
            'skills': {'content': '', 'lines': [], 'blocks': []},
            'education': {'content': '', 'lines': [], 'blocks': []},
            'projects': {'content': '', 'lines': [], 'blocks': []},
            'certifications': {'content': '', 'lines': [], 'blocks': []},
            'awards': {'content': '', 'lines': [], 'blocks': []},
            'publications': {'content': '', 'lines': [], 'blocks': []},
            'volunteer': {'content': '', 'lines': [], 'blocks': []},
            'languages': {'content': '', 'lines': [], 'blocks': []},
            'interests': {'content': '', 'lines': [], 'blocks': []},
            'other': {'content': '', 'lines': [], 'blocks': []}
        }

        current_section = 'header'
        current_job = None

        for i, line in enumerate(lines):
            stripped = line.strip()

            if not stripped:
                continue

            # Detect section headers
            if self._is_section_header_strict(stripped):
                # Save current job if in experience
                if current_job and current_section == 'experience':
                    sections['experience']['jobs'].append(current_job)
                    current_job = None

                # Determine which section
                line_lower = stripped.lower()

                if any(kw in line_lower for kw in ['summary', 'profile', 'objective', 'about', 'overview']):
                    current_section = 'summary'
                elif any(kw in line_lower for kw in ['experience', 'work', 'employment', 'history', 'career']):
                    current_section = 'experience'
                elif any(kw in line_lower for kw in ['skill', 'technical', 'competenc', 'technolog', 'expertise', 'proficienc']):
                    current_section = 'skills'
                elif any(kw in line_lower for kw in ['education', 'academic', 'qualification', 'degree', 'training']):
                    current_section = 'education'
                elif any(kw in line_lower for kw in ['project']):
                    current_section = 'projects'
                elif any(kw in line_lower for kw in ['certification', 'certificate', 'license', 'credential']):
                    current_section = 'certifications'
                elif any(kw in line_lower for kw in ['award', 'honor', 'achievement', 'recognition', 'accomplishment']):
                    current_section = 'awards'
                elif any(kw in line_lower for kw in ['publication', 'research', 'paper']):
                    current_section = 'publications'
                elif any(kw in line_lower for kw in ['volunteer', 'community']):
                    current_section = 'volunteer'
                elif any(kw in line_lower for kw in ['language']):
                    current_section = 'languages'
                elif any(kw in line_lower for kw in ['interest', 'hobbies', 'activities']):
                    current_section = 'interests'
                else:
                    current_section = 'other'

                sections[current_section]['lines'].append(line)
                continue

            # Add to current section
            sections[current_section]['lines'].append(line)

            # Experience section - special handling for jobs
            if current_section == 'experience':
                if self._is_job_title_line(stripped, lines, i):
                    if current_job:
                        sections['experience']['jobs'].append(current_job)

                    current_job = {
                        'title': stripped,
                        'lines': [line],
                        'bullets': []
                    }
                elif current_job:
                    current_job['lines'].append(line)
                    if self._is_bullet_start(stripped):
                        current_job['bullets'].append(stripped)

            # All other sections - each non-empty line is a block
            elif current_section != 'header':
                if stripped and not self._is_section_header_strict(stripped):
                    # For education, check if it looks like a degree line that should be grouped
                    if current_section == 'education':
                        # If this looks like a university/college line and we have blocks
                        if sections['education']['blocks'] and self._looks_like_university(stripped):
                            # Add to previous block instead of creating new one
                            prev_block = sections['education']['blocks'][-1]
                            prev_block['text'] += '\n' + stripped
                        else:
                            sections[current_section]['blocks'].append({
                                'text': stripped,
                                'type': 'bullet' if self._is_bullet_start(stripped) else 'text'
                            })
                    else:
                        sections[current_section]['blocks'].append({
                            'text': stripped,
                            'type': 'bullet' if self._is_bullet_start(stripped) else 'text'
                        })

        # Save last job
        if current_job:
            sections['experience']['jobs'].append(current_job)

        # Build content strings
        for section_name, section_data in sections.items():
            section_data['content'] = '\n'.join(section_data['lines'])

        return sections

    def _is_job_title_line(self, line: str, all_lines: List[str], idx: int) -> bool:
        """Detect job title"""
        date_pattern = r'\b(19|20)\d{2}\b|present|current|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec'

        has_date = bool(re.search(date_pattern, line, re.IGNORECASE))

        if not has_date and idx + 1 < len(all_lines):
            next_line = all_lines[idx + 1].strip()
            has_date = bool(re.search(date_pattern, next_line, re.IGNORECASE))

        title_keywords = ['engineer', 'developer', 'manager', 'analyst', 'designer',
                         'consultant', 'specialist', 'lead', 'senior', 'junior', 'director',
                         'coordinator', 'associate', 'intern', 'architect', 'scientist']

        has_title_keyword = any(kw in line.lower() for kw in title_keywords)
        is_bullet = self._is_bullet_start(line)

        return (has_date or has_title_keyword) and not is_bullet

    def _looks_like_university(self, line: str) -> bool:
        """Check if line looks like a university/college name"""
        university_keywords = [
            'university', 'college', 'institute', 'school', 'academy',
            'polytechnic', 'campus', 'university of', 'state university',
            'tech', 'institute of technology'
        ]

        line_lower = line.lower()

        # Check for university keywords
        has_university = any(kw in line_lower for kw in university_keywords)

        # Check if line has dates (if so, it's more likely a separate entry)
        has_date = bool(re.search(r'\b(19|20)\d{2}\b', line))

        # If has university keyword but no date, likely continuation
        # If has date, it's a separate degree
        if has_university and not has_date:
            return True

        # Check for location patterns (City, State or City, Country)
        has_location = bool(re.search(r',\s*[A-Z]{2}\b|,\s*[A-Z][a-z]+\s*$', line))

        return has_university and has_location

    def extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact info"""
        contact = {
            'name': '',
            'email': '',
            'phone': '',
            'linkedin': '',
            'location': ''
        }

        lines = text.split('\n')[:10]

        for line in lines:
            email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', line)
            if email_match and not contact['email']:
                contact['email'] = email_match.group()

            phone_match = re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', line)
            if phone_match and not contact['phone']:
                contact['phone'] = phone_match.group()

            linkedin_match = re.search(r'linkedin\.com/in/[\w-]+', line, re.IGNORECASE)
            if linkedin_match and not contact['linkedin']:
                contact['linkedin'] = linkedin_match.group()

            if not contact['name'] and line.strip() and '@' not in line and not re.search(r'\d{3}[-.]?\d{3}[-.]?\d{4}', line):
                if len(line.split()) <= 5 and not line.isupper():
                    contact['name'] = line.strip()

        return contact